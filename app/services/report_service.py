import os
import json
import google.generativeai as genai
from flask import current_app
from docx import Document
from datetime import datetime

class ReportService:
    def __init__(self):
        self.data_dir = current_app.config['DATA_DIR']
        self.templates_dir = current_app.config['TEMPLATES_DIR']
        
        # Configure Gemini API if API key is provided; otherwise operate in local-fallback mode
        gemini_key = os.getenv('GEMINI_API_KEY')
        if gemini_key:
            try:
                genai.configure(api_key=gemini_key)
                self.model = genai.GenerativeModel('gemini-2.5-pro')
            except Exception as e:
                # If model initialization fails, fall back to local generator and log
                print(f"[ReportService] Gemini init failed: {e}")
                self.model = None
        else:
            print("[ReportService] GEMINI_API_KEY not set — using local report generator fallback.")
            self.model = None
    
    def generate_report(self, scan_id, report_type):
        """Generate report using Gemini with appropriate template"""
        try:
            print(f"[ReportService] generate_report called with scan_id={scan_id}, report_type={report_type}")
            # Load scan results
            scan_results = self._load_scan_results(scan_id)
            print(f"[ReportService] Loaded scan results for {scan_id}; findings={len(scan_results.get('findings', []))}")
            
            # Load appropriate template
            template = self._load_template(report_type)
            print(f"[ReportService] Loaded template for {report_type} (length={len(template)} chars)")
            
            # Generate report content using Gemini
            report_content = self._generate_with_gemini(scan_results, template, report_type)
            
            # Format and save report
            report_path = self._format_and_save_report(report_content, scan_id, report_type)
            
            return report_path
            
        except Exception as e:
            raise Exception(f"Report generation failed: {str(e)}")
    
    def _load_template(self, report_type):
        """Load template based on report type"""
        template_mapping = {
            'regulatory_compliance': 'regulatory_compliance_template.md',
            'technical_operational': 'technical_operational_template.md',
            'business_focused': 'business_focused_template.md'
        }
        
        template_file = template_mapping.get(report_type)
        if not template_file:
            raise ValueError(f"Unknown report type: {report_type}")
        
        template_path = os.path.join(self.templates_dir, template_file)
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")
        
        with open(template_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _generate_with_gemini(self, scan_results, template, report_type):
        """Use Gemini to generate report content"""
        prompt = f"""
        You are a professional security report generator. Use the provided template and security scan data to create a comprehensive {report_type.replace('_', ' ')} report.

        TEMPLATE:
        {template}

        SECURITY SCAN DATA:
        {json.dumps(scan_results, indent=2)}

        Instructions:
        1. Follow the template structure exactly
        2. Replace placeholders with actual data from the scan results
        3. Write in professional, clear language
        4. Include specific findings with file paths and line numbers where applicable
        5. Provide actionable recommendations
        6. Ensure compliance mappings are accurate
        
        Generate the complete report:
        """
        
        # If the model is not available, use a local generator fallback
        if not self.model:
            return self._local_generate(scan_results, template, report_type)

        try:
            response = self.model.generate_content(prompt)
            # Some genai responses may wrap content differently; attempt to return the textual output
            return getattr(response, 'text', str(response))
        except Exception as e:
            print(f"[ReportService] Gemini generation failed: {e}")
            # Fallback to local generation rather than raising to avoid 500s
            return self._local_generate(scan_results, template, report_type)

    def _local_generate(self, scan_results, template, report_type):
        """A minimal local report generator used when Gemini is unavailable.
        It merges the template with a findings summary and repository context.
        """
        parts = []

        # Header and template inclusion
        parts.append(f"# {report_type.replace('_', ' ').title()} (Local Generated)")
        parts.append('\n')
        parts.append(template)

        # Executive summary derived from repo context (README) when available
        repo_info = scan_results.get('repo_info', {})
        parts.append('\n\n## Executive Summary')
        readme_text = None
        for k, v in repo_info.get('found', {}).items():
            if k.startswith('README'):
                readme_text = v.get('content')
                break

        if readme_text:
            # include a short excerpt from README to give business context
            excerpt = readme_text.strip()[:1200]
            parts.append('Context from repository README:')
            parts.append(excerpt + ('...' if len(readme_text) > 1200 else ''))
        else:
            parts.append('No README found; executive summary is derived from scan findings and templates.')

        # Scan Summary
        parts.append('\n\n## Scan Summary')
        summary = scan_results.get('summary', {})
        parts.append(f"- Total findings: {summary.get('total_findings', 0)}")
        severity = summary.get('severity_breakdown', {})
        for sev, count in severity.items():
            parts.append(f"- {sev}: {count}")

        # Detailed Findings / Analysis
        parts.append('\n\n## Detailed Findings and Analysis')
        findings = scan_results.get('findings', [])
        if not findings:
            parts.append('No findings detected during automated analysis.')
        else:
            for idx, f in enumerate(findings, 1):
                title = f.get('title') or f.get('shortform_keyword')
                short = f.get('shortform_keyword')
                path = f.get('file_path') or 'N/A'
                line = f.get('line_number')
                sev = f.get('severity')
                desc = f.get('description') or ''
                snippet = f.get('context_snippet') or ''
                remediation = f.get('remediation') or 'Refer to template remediation.'
                compliance = f.get('compliance') or []

                parts.append(f"### {idx}. {title} ({short})")
                parts.append(f"- Severity: {sev}")
                parts.append(f"- Location: {path}:{line}")
                if desc:
                    parts.append(f"- Description: {desc}")
                if snippet:
                    parts.append('\n**Evidence (code snippet):**')
                    # keep snippet short
                    parts.append('```')
                    parts.append(snippet.strip()[:800])
                    parts.append('```')
                parts.append(f"- Recommendation: {remediation}")
                if compliance:
                    parts.append(f"- Compliance mappings: {', '.join(compliance)}")
                parts.append('\n')

        # Gemini Analysis Integration
        gemini_analysis = scan_results.get('gemini_analysis', {})
        if gemini_analysis:
            parts.append('\n\n## Gemini Analysis')
            for section, content in gemini_analysis.items():
                parts.append(f"### {section}")
                parts.append(content)

        # Repository Policy and Context excerpts
        if repo_info.get('policy_files'):
            parts.append('\n\n## Repository Policies (excerpts)')
            for ppath, content in repo_info.get('policy_files', {}).items():
                excerpt = content.strip()[:1000]
                parts.append(f"### {ppath}")
                parts.append(excerpt + ('...' if len(content) > 1000 else ''))

        # Methodology and Limitations
        parts.append('\n\n## Methodology')
        parts.append('Automated scanners used:')
        parts.append('- CodeT5-based static heuristics (pattern checks)')
        parts.append('- SCA heuristics scanning requirements.txt and package.json')
        parts.append('- Repository context extractor for README and policy files')

        parts.append('\n\n## Limitations')
        parts.append('- The CodeT5 analyzer currently only processes Python files; other languages may not be covered.')
        parts.append('- SCA is heuristic-based and does not perform CVE lookups or transitive analysis.')
        parts.append('- This automated report is intended as a starting point; manual review is recommended for high-severity findings.')

        parts.append('\n\n-- End of report (generated locally) --')
        return '\n\n'.join(parts)
    
    def _load_scan_results(self, scan_id):
        """Load scan results from storage"""
        results_path = os.path.join(self.data_dir, 'scanned_results', f'{scan_id}.json')
        
        if not os.path.exists(results_path):
            raise FileNotFoundError(f"Scan results not found for ID: {scan_id}")
        
        with open(results_path, 'r') as f:
            return json.load(f)
    
    def _format_and_save_report(self, content, scan_id, report_type):
        """Format report content and save as document"""
        # Create reports directory if it doesn't exist
        reports_dir = os.path.join(self.data_dir, 'generated_reports')
        os.makedirs(reports_dir, exist_ok=True)
        
        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{report_type}_{scan_id}_{timestamp}.docx"
        report_path = os.path.join(reports_dir, filename)
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Security {report_type.replace("_", " ").title()} Report', 0)
        
        # Split content into paragraphs and add to document
        paragraphs = content.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                if paragraph.startswith('#'):
                    # Handle headers
                    level = paragraph.count('#')
                    text = paragraph.strip('#').strip()
                    doc.add_heading(text, level)
                else:
                    doc.add_paragraph(paragraph.strip())
        
        # Save document
        doc.save(report_path)
        return report_path
